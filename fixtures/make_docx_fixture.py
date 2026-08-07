"""One-off script to generate the DOCX evidence fixture. Run once after deps are installed.

SYNTHETIC TEST FIXTURE — the content this script generates is invented for this
project's tests. It does not describe any real organization's policies.
"""

from pathlib import Path

import docx

doc = docx.Document()

doc.add_paragraph(
    "SYNTHETIC TEST FIXTURE — NOT A REAL DOCUMENT. This file was invented for this "
    "project's test fixtures. It does not describe the policies of any real "
    "organization and must not be treated as evidence about any company's actual "
    "security controls."
)

doc.add_heading("Business Continuity Plan", level=1)

doc.add_heading("Backup Strategy", level=2)
doc.add_paragraph(
    "Production databases are backed up hourly with 30-day retention. Backups are "
    "stored in a separate cloud region from the primary production environment."
)

doc.add_heading("Disaster Recovery", level=2)
doc.add_paragraph(
    "In the event of a regional outage, services fail over to a warm standby "
    "environment in a second region. The target recovery time objective (RTO) is 4 "
    "hours and the target recovery point objective (RPO) is 1 hour."
)

doc.add_heading("Incident Response", level=1)
doc.add_heading("Escalation", level=2)
doc.add_paragraph(
    "Security incidents are triaged by the on-call engineer within 15 minutes of "
    "detection and escalated to the security lead if classified as high severity."
)

out_path = Path(__file__).parent / "evidence" / "business_continuity_plan.docx"
doc.save(str(out_path))
print(f"Wrote {out_path}")
